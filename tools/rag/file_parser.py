from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


def _normalize_text(text: str) -> str:
    lines = [line.rstrip() for line in str(text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if line.strip():
            blank_run = 0
            cleaned.append(line)
            continue
        blank_run += 1
        if blank_run <= 2:
            cleaned.append("")
    return "\n".join(cleaned).strip()


@dataclass(frozen=True)
class ParsedKnowledgeDocument:
    text: str
    parser_name: str
    metadata: dict[str, Any] = field(default_factory=dict)


def _parse_txt(data: bytes) -> ParsedKnowledgeDocument:
    text = data.decode("utf-8", errors="replace")
    return ParsedKnowledgeDocument(text=_normalize_text(text), parser_name="txt")


def _parse_md(data: bytes) -> ParsedKnowledgeDocument:
    text = data.decode("utf-8", errors="replace")
    return ParsedKnowledgeDocument(text=_normalize_text(text), parser_name="markdown")


def _parse_docx(data: bytes) -> ParsedKnowledgeDocument:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on optional env state
        raise RuntimeError("python-docx is required to parse DOCX knowledge files") from exc

    document = Document(io.BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    text = "\n\n".join(paragraphs)
    return ParsedKnowledgeDocument(
        text=_normalize_text(text),
        parser_name="docx",
        metadata={"paragraph_count": len(paragraphs)},
    )


def _parse_pdf(data: bytes) -> ParsedKnowledgeDocument:
    try:
        import fitz
    except ImportError as exc:  # pragma: no cover - depends on optional env state
        raise RuntimeError("pymupdf is required to parse PDF knowledge files") from exc

    doc = fitz.open(stream=data, filetype="pdf")
    page_texts: list[str] = []
    for page in doc:
        text = str(page.get_text("text") or "").strip()
        if text:
            page_texts.append(text)
    text = "\n\n".join(page_texts)
    return ParsedKnowledgeDocument(
        text=_normalize_text(text),
        parser_name="pdf",
        metadata={"page_count": len(page_texts)},
    )


def parse_uploaded_file(data: bytes, *, filename: str, content_type: str = "") -> ParsedKnowledgeDocument:
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix == "pdf":
        return _parse_pdf(data)
    if suffix == "docx":
        return _parse_docx(data)
    if suffix == "md":
        return _parse_md(data)
    if suffix == "txt":
        return _parse_txt(data)
    raise ValueError(f"Unsupported knowledge file type: {suffix or content_type or 'unknown'}")


__all__ = ["ParsedKnowledgeDocument", "parse_uploaded_file"]
